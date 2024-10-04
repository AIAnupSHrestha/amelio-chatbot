from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime


def DocumentHeader(doc, BrandColor, BrandLogo):
    header = doc.sections[0].header
    paragraph = header.paragraphs[0]
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    
    # Create a border (colored line) element
    borders = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:top')
    bottom.set(qn('w:val'), 'single')  # Line style
    bottom.set(qn('w:sz'), '150')  # Thickness (1/8 of a point)
    bottom.set(qn('w:space'), '1')  # Space between text and line
    bottom.set(qn('w:color'), BrandColor)  # Company Hex color code
    borders.append(bottom)
    pPr.append(borders)

    logo = header.paragraphs[0]
    run = logo.add_run("\n")
    run.add_picture(BrandLogo, width=Inches(0.8), height=Inches(0.5))

    return "Header Created !!"

def PageOne(doc, CompanyName, CompanyAddress):
    current_datetime = datetime.now()
    current_date = current_datetime.date()

    for _ in range(7):  # empty lines
        doc.add_paragraph("")
    # Adding Current Date    
    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format = date.paragraph_format
    format.left_indent = Inches(1)
    format.right_indent = Inches(1)
    run = date.add_run(f"{current_date}")
    run.bold = True
    run.font.name = 'Open Sans'
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(96,96,96)

    # Adding Title of the Document
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format = title.paragraph_format
    format.left_indent = Inches(1)
    format.right_indent = Inches(1)
    run = title.add_run("Politique d'aménagement du temps de travail")
    run.bold = True
    run.font.name = 'Open Sans'  
    run.font.size = Pt(32)
    run.font.color.rgb = RGBColor(96,96,96)

    for _ in range(9):  # empty lines
        doc.add_paragraph("")

    # Adding company details
    companydetails = doc.add_paragraph()
    companydetails.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    format = companydetails.paragraph_format
    format.left_indent = Inches(1)
    format.right_indent = Inches(1)
    run = companydetails.add_run(CompanyName)
    run.bold = True
    run.font.name = 'Open Sans'
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(96,96,96)

    run = companydetails.add_run("\n")

    run = companydetails.add_run(CompanyAddress)
    run.bold = True
    run.font.name = 'Open Sans'
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(96,96,96)

    print("ForntPage Created")


def QuestionAnswer(doc): #, flexible_response, applied_context_response, eligibility_criteria_response, missing_response):
    footer = doc.sections[0].footer
    footer_paragraph = footer.paragraphs[0]

    # Create an OxmlElement for the page number field
    field_code = OxmlElement('w:fldSimple')
    field_code.set(qn('w:instr'), 'PAGE')
    run = OxmlElement('w:r')
    text = OxmlElement('w:t')
    text.text = ' '
    run.append(text)
    field_code.append(run)
    footer_paragraph._element.append(field_code)

    footer_paragraph.alignment = 1

    flexible_response = ["response1", "response2", "response3"] #tracker.get_slot("user_response")
    applied_context_response = ["response1", "response2", "response3"] #tracker.get_slot("applied_context_response")
    eligibility_criteria_response = ["response1", "response2", "response3"] #tracker.get_slot("eligibility_criteria_response")
    missing_response = ["response1", "response2", "response3"] #tracker.get_slot("missing_element_response")
    flexible_question = ["Question1", "Question2", "Question3"] #flexible_work_questions[0]
    applied_question = ["Question1", "Question2", "Question3"] #applied_content_questions[0]
    eligibility_question = ["Question1", "Question2", "Question3"] #eligibility_criteria_question[0]
    missing_question =  ["Question1", "Question2", "Question3"] #missing_element_questions[0]
    title = doc.add_heading('HR Policy', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    run.font.name = 'Open Sans'
    run.font.color.rgb = RGBColor(0,0,0)
       
    policy = doc.add_heading('Policy Question Answers', level=2)
    format = policy.paragraph_format
    format.left_indent = Inches(1)
    format.right_indent = Inches(1)
    policy.add_run("\n")
    run = policy.runs[0]
    run.font.color.rgb = RGBColor(0,0,0)
    run.font.name = 'Open Sans'
    run.font.size = Pt(16)
    policy.add_run("\n")

    flexiblePolicy = doc.add_paragraph()
    format = flexiblePolicy.paragraph_format
    format.left_indent = Inches(1)
    format.right_indent = Inches(1)
    run = flexiblePolicy.add_run("Flexible work policy")
    run.bold = True
    flexiblePolicy.add_run("\n")
    for num in range(len(flexible_question)):
        run = flexiblePolicy.add_run(flexible_question[num])
        flexiblePolicy.add_run("\n")
        run = flexiblePolicy.add_run(f"Ans: {flexible_response[num]}")
        flexiblePolicy.add_run("\n")

    for run in flexiblePolicy.runs:
        run.font.name = 'Open Sans'
        run.font.size = Pt(10)

    appliedContext = doc.add_paragraph()
    format = appliedContext.paragraph_format
    format.left_indent = Inches(1)
    format.right_indent = Inches(1)
    run = appliedContext.add_run("Applied Context")
    run.bold = True
    run.font.name = 'Open Sans'
    appliedContext.add_run("\n")
    for num in range(len(applied_question)):
        run = appliedContext.add_run(applied_question[num])
        appliedContext.add_run("\n")
        run = appliedContext.add_run(f"Ans: {applied_context_response[num]}")
        appliedContext.add_run("\n")

    for run in appliedContext.runs:
        run.font.name = 'Open Sans'
        run.font.size = Pt(10)

    eligibilityCriteria = doc.add_paragraph()
    format = eligibilityCriteria.paragraph_format
    format.left_indent = Inches(1)
    format.right_indent = Inches(1)
    run = eligibilityCriteria.add_run("Eligibility Criteria")
    run.bold = True
    run.font.name = 'Open Sans'
    eligibilityCriteria.add_run("\n")
    for num in range(len(eligibility_question)):
        run = eligibilityCriteria.add_run(eligibility_question[num])
        eligibilityCriteria.add_run("\n")
        run = eligibilityCriteria.add_run(f"Ans: {eligibility_criteria_response[num]}")
        eligibilityCriteria.add_run("\n")

    for run in eligibilityCriteria.runs:
        run.font.name = 'Open Sans'
        run.font.size = Pt(10)

    missingElement = doc.add_paragraph()
    format = missingElement.paragraph_format
    format.left_indent = Inches(1)
    format.right_indent = Inches(1)
    run = missingElement.add_run("Missing Element")
    run.bold = True
    run.font.name = 'Open Sans'
    missingElement.add_run("\n")
    for num in range(len(missing_question)):
        run = missingElement.add_run(missing_question[num])
        missingElement.add_run("\n")
        run = missingElement.add_run(f"Ans: {missing_response[num]}")
        missingElement.add_run("\n")

    for run in missingElement.runs:
        run.font.name = 'Open Sans'
        run.font.size = Pt(10)

    return ("Question Answered Created !!")    


def createDocument(): #company_name, brand_color,brand_logo, flexible_response, applied_response, eligibility_response, missing_response):
    

    company_name = "CompanyName"
    doc = Document("OpenAI - Document generator-Test project.docx")
    flexible_response = ["response1", "response2", "response3"]
    # ForntPage = doc.sections[0]
    # ForntPage.top_margin = Inches(0)  
    # ForntPage.header_distance = Inches(0)
    # ForntPage.left_margin = 0
    # ForntPage.right_margin = 0
    # print(DocumentHeader(doc, '990000', 'logo-search-grid-1x.jpg'))
    # print(PageOne(doc, "CompanyName", "CompanyAddress"))

    # # ForntPage.footer.is_linked_to_previous = False
    # # ForntPage.different_first_page_header_footer = True   
    
    # Section1 = doc.add_section(start_type=2)
    # Section1.left_margin = Inches(0)
    # Section1.right_margin = Inches(0)
    # print(QuestionAnswer(doc))

    # text_to_replace = "Everyone's responsibilities"
    # for paragraph in doc.paragraphs:
    #     if paragraph.text == text_to_replace:
    #         for run in paragraph.runs:
    #             if text_to_replace in run.text:
    #                 run.text = run.text.replace(old_text, new_text)
        
        # If you want to get more granular control, iterate through runs in each paragraph
        
            # for run in paragraph.runs:
                # print("Run text:", run.text)

        # for item in bullet_points:
        #     paragraph = doc.add_paragraph(item, style='ListBullet')



    for section in doc.sections:
        if section.start_type == 2: 
            # print(f"Section start type: {section.start_type}")
            # print(f"Page height: {section.page_height}")
            # print(f"Page width: {section.page_width}")
            # print(f"Orientation: {section.orientation}")
            
            # You can also access the content of paragraphs within the section
            # print("Section content:")
            for paragraph in doc.paragraphs:
                text_to_replace = "Everyone's responsibilities"
                if paragraph.text == text_to_replace:
                    print(f"Paragraph text: {paragraph.text}")
                    for run in paragraph.runs:
                        if text_to_replace in run.text:
                            run.text = run.text.replace(text_to_replace, "Flexible Work Policy")
                for item in flexible_response:
                    p = doc.add_paragraph('•  ' + item)
                    p.paragraph_format.left_indent = Inches(0.5)
                    # p.add_run(f"- {item}")#.bold = True
                    if "Immediate supervisor:" in paragraph.text:
                        # Insert the list paragraph before it
                        # Move the cursor to the paragraph before the target
                        doc._element.body.insert(doc._element.body.index(paragraph._element), p._element)

    doc.save(f'{company_name}_HR_document.docx')
    # return "Document Created..."

# company_name = "test Company"
# brand_color = "#FF0000"
# brand_logo = 'pexels-helloaesthe-28056131.jpg'
# tracker = []
# flexible_response = tracker.get_slot("user_response")
# applied_context_response = tracker.get_slot("applied_context_response")
# eligibility_criteria_response = tracker.get_slot("eligibility_criteria_response")
# missing_response = tracker.get_slot("missing_element_response")
document = createDocument() #(company_name, brand_color,brand_logo, flexible_response, applied_response, eligibility_response, missing_response)
# print(document)



































































































































































































































































































































































































# import cohere
# import os
# from dotenv import load_dotenv
# import PyPDF2

# load_dotenv()

# COHERE_KEY = os.getenv("COHERE_API_KEY")

# co = cohere.Client(api_key=COHERE_KEY)


# def prompt_engineering(prompt):
#     response = co.chat(
#         message=prompt,
#         model="command-r-plus-08-2024",
#         temperature=0.3
#     )
#     return response.text

# ai_observation_prompt_template = """
# This is a remote work HR policy, are we missing something in this HR policy?
# Identify the missing elements in the following remote work HR policy:

# {policy}

# If there are any missing element in the provided policy, Return the missing elements separated by commas.
# """

# missing_element_prompt = """
# "Generate questions that address the following missing elements in a Remote Work HR policy: {missing_elements}
# Avoid compound and tail questions under any circumstances. 
# Do not provide any additional information beyond the questions"

# Separator: |
# """


# """
# Given the following list missing elements in the HR policy: 

# {missing_elements}

# Generate at two separate questions for each element, ensuring the questions are separated by a pipe character (|) only. 
# Avoid compound or tail questions under any circumstances. 
# Do not provide any additional information beyond the questions and do not add missing element in your response.
# """

# def extract_text_from_pdf(pdf_file_path):
#         try:
#             with open(pdf_file_path, 'rb') as pdf_file:
#                 pdf_reader = PyPDF2.PdfReader(pdf_file)
#                 text = ''
#                 for page_num in range(len(pdf_reader.pages)):
#                     page = pdf_reader.pages[page_num]
#                     text += page.extract_text()
#                 return text
#         except Exception as e:
#             print(f"Error extracting text from PDF: {e}")
#             return None

# missing_elements = []
# questions = []

# def generate(pdf_file_path):
#         extracted_text = extract_text_from_pdf(pdf_file_path)
#         print("extracted!!")
#         prompt = ai_observation_prompt_template.format(policy = extracted_text)
#         missing_element = prompt_engineering(prompt=prompt)
#         ele = missing_element.replace(", and", ",")
#         elements_list = [element.strip() for element in ele.split(',') if element.strip()]
#         missing_elements.append(elements_list)
#         print("Generated missing ele!!")
#         prompt = missing_element_prompt.format(missing_elements = missing_elements[0])
#         question = prompt_engineering(prompt=prompt).replace("- ", "")
#         question = question.replace("\n", "")

#         question_list = [que.strip() for que in question.split('|')]
#         questions.append(question_list)
#         print("Generated questions!!")



# pdf_file_path = "Policy_Questions_and_Responses.pdf"
# generate(pdf_file_path)
# print(missing_elements[0])
# for question in questions[0]: 
#      print(question)
