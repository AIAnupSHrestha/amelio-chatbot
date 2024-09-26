# rasa run --enable-api --cors="*"
# rasa run actions

# This files contains your custom actions which can be used to run
# custom Python code.
#
# See this guide on how to implement these action:
# https://rasa.com/docs/rasa/custom-actions


# This is a simple example for a custom action which utters "Hello World!"

from typing import Any, Dict, List, Text, Union
import random
from rasa_sdk import Action, Tracker
from rasa_sdk.executor import CollectingDispatcher
from rasa_sdk.events import SlotSet, Form, FollowupAction
from rasa_sdk import Action, FormValidationAction, forms

import json

import os
from openai import OpenAI
from dotenv import load_dotenv

# import re
from fpdf import FPDF
import PyPDF2

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import ollama
from reportlab.pdfgen import canvas 
from reportlab.pdfbase.ttfonts import TTFont 
from reportlab.pdfbase import pdfmetrics 
from reportlab.lib import colors 
import mysql.connector

import cohere


# import mysql.connector
# from mysql.connector import errorcode

load_dotenv()

OPENAI_KEY = "12345678765432" #os.getenv("OPENAI_API_KEY")
client = OpenAI(
  api_key=OPENAI_KEY
)

COHERE_KEY = os.getenv("COHERE_API_KEY")

co = cohere.Client(api_key=COHERE_KEY)

    # response = client.chat.completions.create(
    #     model="gpt-3.5-turbo-0125",
    #     messages=[
    #         {"role": "user", "content": prompt},
    #     ],
    # )
    # return response.choices[0].message.content
def prompt_engineering(prompt):
    response = co.chat(
        message=prompt,
        model="command-r-plus-08-2024",
        temperature=0.3
    )
    return response.text




# config = {
#     'host': '127.0.0.1',  # or your MySQL server host
#     'user':'root', 
#     'password':"",
#     # Add 'user' and 'password' fields if your MySQL server requires them
# }
# cnx = mysql.connector.connect(**config)
# cursor = cnx.cursor()
# cursor.execute("USE amelio")

# PROMPT_TEMPLATE = """
# You are an expert in drafting HR policies. I am currently working on creating an {job_type}. The policy will include the following condition: {flexible_work_option}.
# Based on this specific clause, please provide a Python list of detailed questions to consider.
# The output should be a Python list in the following format:
# [
#     "Question 1",
#     "Question 2",
#     "Question 3",
#     "Question 4",
#     "Question 5"
# ]
# """

PROMPT_TEMPLATE = """
You are an expert in drafting HR policies. I am currently working on creating a {job_type}. 
The policy will include the following condition: {flexible_work_option}. 
Based on this specific clause, please provide 3 detailed questions to consider, separated by pipe character (|) only.
And avoid generating tail and compound questions in any circumtances.
"""

with open('actions/predefined_questions.json', 'r') as file:
    predefined_questions = json.load(file)

def get_policy_from_db(table_name, attribute_name, value):
    value = value.lower()
    query = f"SELECT * FROM {table_name} WHERE {attribute_name}='{value}'"
    # cursor.execute(query)
    # result = cursor.fetchmany()
    # Returns list of policy
    # return result



class ActionGreet(Action):

    def name(self) -> Text:
        return "action_greet"

    def run(self, dispatcher: CollectingDispatcher,
            tracker: Tracker,
            domain: Dict[Text, Any]) -> List[Dict[Text, Any]]:

        messages = ["Hi there. 👋😃 It's such a pleasure to have you here. How may I help you?",
                    "Hello, 🤗 how can we assist you?"]
        
        buttons = [
            # {"payload": "/job_posting", "title": "Job Posting"},
            # {"payload": '/hr_policy', "title": "HR Policy"},
            {"payload": '/select_option{"option": "selected_hr_policy"}', "title": "HR Policy"}
        ]

        reply = random.choice(messages)
        dispatcher.utter_message(text=reply, buttons=buttons)
        return []
    
class ActionCompanyName(Action):
    def name(self) -> Text:
        return "action_company_name"
    
    def run(self, dispatcher: CollectingDispatcher,
                tracker: Tracker,
                domain: Dict[Text, Any]) -> List[Dict[Text, Any]]:
        dispatcher.utter_message(text="You have selected HR policy. \nWhat is the name of your company?")
        return[SlotSet("indicator", "company_name")]
    
class ActionSetCompanyName(Action):
    def name(self) -> Text:
        return "action_set_company_name"
    
    def run(self, dispatcher: CollectingDispatcher,
                tracker: Tracker,
                domain: Dict[Text, Any]) -> List[Dict[Text, Any]]:
        company_name = tracker.latest_message.get("text")
        return[SlotSet("company_name", company_name),
               FollowupAction("action_hr_policy")]


class ActionHrPolicy(Action):
    
        def name(self) -> Text:
            return "action_hr_policy"
    
        def run(self, dispatcher: CollectingDispatcher,
                tracker: Tracker,
                domain: Dict[Text, Any]) -> List[Dict[Text, Any]]:
            
            # buttons = [
            #     {"payload": "/hr_policy{'content_type': 'hr_policy'}", "title": "HR Policy"},
            #     {"payload": "/job_posting", "title": "Job Posting"}
            # ]
            try:
                hr_policy_type = tracker.get_slot("hr_policy_type")
            except:
                hr_policy_type = None
    
            if hr_policy_type:
                message = f"You selected: {hr_policy_type} policy."
                # policies = get_policy_from_db('hr_policy', 'policy_name', hr_policy_type)
                # Fetch all of remote work policies
                # query = f"SELECT hr_policy_type.* \
                #             FROM hr_policy_type \
                #             JOIN hr_policy ON hr_policy_type.hr_policy_id = hr_policy.id \
                #             WHERE hr_policy.policy_name = '{hr_policy_type}'; \
                #         "
                # cursor.execute(query)
                # policies = cursor.fetchmany()
                policies = [
                    ('flexible', 'Flexible Work'),
                    ('remote', 'Remote Work'),
                    ('part_time', 'Part-time Work'),
                    ('unpaid_leave', 'Unpaid Leave'),
                    ('job_sharing', 'Job Sharing')
                ]
                buttons = []
                if policies:
                    message += "\nHere are some templates for you:"
                    for p_template in policies:
                        # Suggest as button
                        buttons.append({
                            "payload": '/policy_type{"policy_name": "'+p_template[0]+'"}',
                            "title": p_template[1].capitalize()
                        })
                buttons.append(
                    {"payload": '/policy_type{"policy_name": None}', "title": "Create your own"}
                )
                dispatcher.utter_message(text=message, buttons=buttons)
            else:
                message = "What policy would you like to create?"
                dispatcher.utter_message(text=message,)
    
            return []

class ActionPolicyType(Action):
        
        def name(self) -> Text:
            return "action_policy_type"
    
        def run(self, dispatcher: CollectingDispatcher,
                tracker: Tracker,
                domain: Dict[Text, Any]) -> List[Dict[Text, Any]]:
            try:
                selected_policy = tracker.get_slot('policy_name')
            except:
                selected_policy = None
            # query = f"SELECT * FROM hr_policy_type WHERE {attribute_name}='{value}'"
            # cursor.execute(query)
            # result = cursor.fetchmany()
            # selected_policy = get_policy_from_db('hr_policy_type', 'policy_name', selected_policy)

            # selected_policy = 'flexible'

            if selected_policy.lower(): # == 'flexible':
                options = predefined_questions[selected_policy]
                buttons = []
                for key, option_val in options.items():
                    # buttons.append(
                    #     {
                    #         "label" : option_val,
                    #         "value": "/select_flexible_work_option{'flexible_work_option': '"+key+"'}"
                    #     }
                    # )
                    # "/select_flexible_work_option{'flexible_work_option': 'a'}"
                # message={"payload":"dropDown","data":buttons}
                    buttons.append({
                        "payload": '/select_flexible_work_option{"flexible_work_option": "'+key+'"}',
                        "title": option_val
                    })
                dispatcher.utter_message(text="Please select your flexible work option:", buttons=buttons)
                # dispatcher.utter_message(text="Please select your flexible work option:", attachment=message)
                # dispatcher.utter_message(text="Please select your flexible work option:", json_message=message)
            else:
                dispatcher.utter_message(text="No policy type selected")
            return []

flexible_work_questions = []

class ActionSelectFlexibleWorkOption(Action):

    def name(self) -> Text:
        return "action_select_flexible_work_option"


    def run(self, dispatcher: CollectingDispatcher, tracker: Tracker, domain: Dict[Text, Any]) -> List[Dict[Text, Any]]:
        flexible_work_option = tracker.get_slot('flexible_work_option')
        option = tracker.get_slot('option').split('_')
        option = ' '.join(option[1:])
        hr_policy_type = tracker.get_slot('hr_policy_type')
        job_type = f'{hr_policy_type} for {option} work'
        selected_policy = tracker.get_slot('policy_name')
        flexible_work_option = predefined_questions[selected_policy].get(flexible_work_option)
        prompt = PROMPT_TEMPLATE.format(job_type=job_type, flexible_work_option=flexible_work_option)
        questions_list = prompt_engineering(prompt=prompt).replace("\n", "")
        flexible_questions = [question.strip() for question in questions_list.split('|')]
        flexible_work_questions.append(flexible_questions)
        # print(flexible_work_questions)

        # for option in flexible_work_option:
        #     option = tracker.get_slot('option').split('_')
        #     option = ' '.join(option[1:])
        #     hr_policy_type = tracker.get_slot('hr_policy_type')
        #     job_type = f'{hr_policy_type} for {option} work'
        #     selected_policy = tracker.get_slot('policy_name')
        #     flexible_work_option = predefined_questions[selected_policy].get(flexible_work_option)
        #     prompt = PROMPT_TEMPLATE.format(job_type=job_type, flexible_work_option=flexible_work_option)
        #     self.generate_questions(prompt)
        indicator = "flexible_work"
        questions = "question" #flexible_questions["question_list"]
        attachments = {
            "questions": questions,
            "payload": "question_list"
        }
        # dispatcher.utter_attachment(attachment=attachments)
        dispatcher.utter_message(text=f"You have selected: {flexible_work_option}")
        dispatcher.utter_message(text="Please answer the following questions:")
        return [SlotSet("indicator", indicator),
                FollowupAction("action_set_question")]
        # return [SlotSet('flexible_work_option', flexible_work_option)]


# class ActionGetQuestions(Action):
#     def name(self) -> Text:
#         return "action_get_questions"

#     def run(self, dispatcher: CollectingDispatcher, tracker: Tracker, domain: Dict[Text, Any]) -> List[Dict[Text, Any]]:

#         questions = flexible_questions["question"]
#         return [{"questions": questions}]




yes_no_prompt = """
    Is the following answer relevant to the question {current_question}: '{user_answer}'? Answer with 'yes' or 'no', without any additional explanation. .
    """  


class ValidateQuestionForm(FormValidationAction):
    def name(self):
        return "action_validate_flexible_work"
    
    def run(self, dispatcher: CollectingDispatcher, tracker: Tracker, domain: Dict[Text, Any]) -> List[Dict[Text, Any]]:
        current_question = tracker.get_slot("question0")
        user_answer = tracker.latest_message.get('text')
        index = int(tracker.get_slot("question_index"))
        response_list = tracker.get_slot("user_response")

        # print(response_list)
        
        if response_list:
            updated_response_list = response_list + [user_answer]
        else:
            updated_response_list = user_answer
        
        prompt = yes_no_prompt.format(current_question=current_question, user_answer=user_answer)

        try:
            answer_relevance = prompt_engineering(prompt=prompt).lower()
        except Exception as e:
            print(f"Error occurred: {e}")

        print(answer_relevance)

        if answer_relevance == "yes.":
            updated_index = index + 1
            print(f"flexible - {updated_response_list}")
            return[SlotSet("question_index", updated_index),
                   SlotSet("user_response", updated_response_list),
                   FollowupAction("action_set_question")]
            # return {"user_response": updated_response_list}
        
        if answer_relevance == "no.":
            dispatcher.utter_message(text="Please enter answer relevant to the question.")
            return[FollowupAction("action_set_question")]
            # return {"question_index": updated_index}
    
    def submit(self, dispatcher, tracker, domain) -> List[Dict]:
        return [FollowupAction("action_set_question")]
    
with open('actions/flexible_questions.json', 'r') as file:
    flexible_questions = json.load(file)
try:
    with open('actions/questions.json', 'r') as file:
        questions = json.load(file)
except Exception as e:
    pass

class ActionSetQuestion(Action):
    def name(self) -> str:
        return "action_set_question"
    

    def run(self, dispatcher: CollectingDispatcher, tracker: Tracker, domain: Dict[Text, Any]) -> List[Dict[Text, Any]]:
        question_list = flexible_work_questions[0]
        print(question_list)

        index = int(tracker.get_slot("question_index"))
        # print(index, question_list[index])
        # print(f"set question index - {index}")

        response_list = tracker.get_slot("user_response")
        print(response_list)

        # question_index = index
        if index >= len(question_list):
            # return [FollowupAction("action_store_response")]
            return [FollowupAction("action_select_applied_context"),
                    SlotSet("question_index", 0)]
        else:
            dispatcher.utter_message(text=question_list[index])
            return[SlotSet("question0", question_list[index])]
            # return [SlotSet("question0", question_list[question_index])]
    
##########################################################################################################
with open('actions/apply_flexible_work_policy_questions.json', 'r') as file:
    applied_questions = json.load(file)
import pprint

class ActionSelectAppliedContexts(Action):
    def name(self) -> Text:
        return "action_select_applied_context"
    
    def run(self, dispatcher: CollectingDispatcher, tracker: Tracker, domain: Dict[Text, Any]) -> List[Dict[Text, Any]]:
        tracker.get_slot("user_response")
        options = applied_questions["applied_context"]
        buttons = []
        for key, option_val in options.items():

            buttons.append({
                    "payload": '/select_applied_context_option{"applied_context_option": "'+key+'"}',
                    "title": option_val
                })
        dispatcher.utter_message(text="Select the work situations where flexible work policy you would like to apply:", buttons=buttons)

        return []

APPLIED_CONTEXT_PROMPT_TEMPLATE ="""
You are an expert in drafting HR policies. I am currently working on creating a {job_type}. 
The policy will include the following condition: {applied_context_option}. 
Based on this specific clause, please provide 2 detailed questions to consider, separated by pipe character (|) only.
And avoid generating tail and compound questions in any circumtances.
"""

applied_content_questions = []

class ActionAppliedContext(Action):
    def name(self):
        return "action_applied_content"
    
    def run(self, dispatcher: CollectingDispatcher, tracker: Tracker, domain: Dict[Text, Any]) -> List[Dict[Text, Any]]:
        option = tracker.get_slot('option').split('_')
        option = ' '.join(option[1:])
        hr_policy_type = tracker.get_slot('hr_policy_type')
        job_type = f'{hr_policy_type} for {option} work'

        applied_context_option = tracker.get_slot('applied_context_option')
        applied_context = applied_questions["applied_context"].get(applied_context_option)

        prompt = APPLIED_CONTEXT_PROMPT_TEMPLATE.format(job_type=job_type, applied_context_option=applied_context)
        questions_list = prompt_engineering(prompt=prompt).replace("\n", "")
        applied_question = [question.strip() for question in questions_list.split('|')]
        applied_content_questions.append(applied_question)


        indicator = "applied_context"

        # prompt = APPLIED_CONTEXT_PROMPT_TEMPLATE.format(applied_context_option=applied_context_option)
        # questions = prompt_engineering(prompt=prompt)

        # for option in flexible_work_option:
        #     prompt = APPLIED_CONTEXT_PROMPT_TEMPLATE.format(job_type=job_type, flexible_work_option=flexible_work_option)
        #     self.generate_questions(prompt)

        questions = "question" #flexible_questions["question_list"]
        attachments = {
            "questions": questions,
            "payload": "question_list"
        }
        # dispatcher.utter_attachment(attachment=attachments)
        dispatcher.utter_message(text=f"You have selected: {applied_context}")
        dispatcher.utter_message(text="Please answer the following questions:")
        return [SlotSet("indicator", indicator),
                FollowupAction("action_set_applied_context_question")]
    
class ActionAppliedContextSetQuestion(Action):
    def name(self) -> str:
        return "action_set_applied_context_question"
    

    def run(self, dispatcher: CollectingDispatcher, tracker: Tracker, domain: Dict[Text, Any]) -> List[Dict[Text, Any]]:
        question_list = applied_content_questions[0]
        print(question_list)

        index = int(tracker.get_slot("question_index"))
        print(f"set question index - {index}")
        # print(tracker.get_slot("applied_context_response"))
        if index >= len(question_list):
            return [FollowupAction("action_select_eligibility_criteria"),
                    SlotSet("question_index", 0)]
        else:
            dispatcher.utter_message(text=question_list[index])
            return[SlotSet("question0", question_list[index])]
    
class ValidateAppliedContext(Action):
    def name(self):
        return "action_validate_applied_context_form"
    
    def run(self, dispatcher: CollectingDispatcher, tracker: Tracker, domain: Dict[Text, Any]) -> List[Dict[Text, Any]]:
        current_question = tracker.get_slot("question0")
        user_answer = tracker.latest_message.get('text')
        index = int(tracker.get_slot("question_index"))
        response_list = tracker.get_slot("applied_context_response")

        if response_list:
            updated_response_list = response_list + [user_answer]
        else:
            updated_response_list = user_answer
        
        prompt = yes_no_prompt.format(current_question=current_question, user_answer=user_answer)
        try:
            answer_relevance = prompt_engineering(prompt=prompt).lower()
        except Exception as e:
            print(f"Error occurred: {e}")


        if answer_relevance == "yes.":
            updated_index = index + 1
            print(f"applied - {updated_response_list}")
            return [SlotSet("applied_context_response", updated_response_list),
                    SlotSet("question_index", updated_index),
                    FollowupAction("action_set_applied_context_question")]
            # return {"applied_context_response": updated_response_list}
        elif answer_relevance == "no.":
            dispatcher.utter_message(text="Please enter answer relevant to the question.")
            return[FollowupAction("action_set_applied_context_question")]
            # return {"question_index": updated_index}
    
####################################################################################################################
with open('actions/eligibility_criteria.json', 'r') as file:
    eligibility_criteria_questions = json.load(file)
    
class ActionSelectEligibilityCriteria(Action):
    def name(self) -> Text:
        return "action_select_eligibility_criteria"
    
    def run(self, dispatcher: CollectingDispatcher, tracker: Tracker, domain: Dict[Text, Any]) -> List[Dict[Text, Any]]:

        options = eligibility_criteria_questions["eligibility_criteria"]
        buttons = []
        for key, option_val in options.items():
            buttons.append({
                    "payload": '/select_eligibility_criteria_option{"eligibility_criteria_option": "'+key+'"}',
                    "title": option_val
                })
        dispatcher.utter_message(text="Select the eligibility criteria would you like to include in your flexible work policy:", buttons=buttons)

        return []
    
ELIGIBILITY_CRITERIA_PROMPT_TEMPLATE = """ 
You are an expert in drafting HR policies. I am currently working on creating a {job_type}. 
The policy will include the following condition: {eligibility_criteria_option}. 
Based on this specific clause, please provide 2 detailed questions to consider, separated by pipe character (|) only.
And avoid generating tail or compound questions in any circumtances
"""

eligibility_criteria_question = []

class ActionEligibilityCriteria(Action):
    def name(self):
        return "action_eligibility_criteria"

    def run(self, dispatcher: CollectingDispatcher, tracker: Tracker, domain: Dict[Text, Any]) -> List[Dict[Text, Any]]:
        option = tracker.get_slot('option').split('_')
        option = ' '.join(option[1:])
        hr_policy_type = tracker.get_slot('hr_policy_type')
        job_type = f'{hr_policy_type} for {option} work'

        eligibility_criteria_option = tracker.get_slot('eligibility_criteria_option')
        eligibility_criteria_option = eligibility_criteria_questions["eligibility_criteria"].get(eligibility_criteria_option)

        prompt = ELIGIBILITY_CRITERIA_PROMPT_TEMPLATE.format(job_type=job_type, eligibility_criteria_option=eligibility_criteria_option)
        questions_list = prompt_engineering(prompt=prompt).replace("\n", "")
        eligibility_questions = [question.strip() for question in questions_list.split('|')]
        eligibility_criteria_question.append(eligibility_questions)

        indicator = "eligibility_criteria"
        # prompt = ELIGIBILITY_CRITERIA_PROMPT_TEMPLATE.format(eligibility_criteria_option=eligibility_criteria_option)
        # questions = prompt_engineering(prompt=prompt)

        # for option in eligibility_criteria_option:
        #     prompt = ELIGIBILITY_CRITERIA_PROMPT_TEMPLATE.format(eligibility_criteria_option=eligibility_criteria_option)
        #     self.generate_questions(prompt)

        questions = "question" #flexible_questions["question_list"]
        attachments = {
            "questions": questions,
            "payload": "question_list"
        }
        # dispatcher.utter_attachment(attachment=attachments)
        dispatcher.utter_message(text=f"You have selected: {eligibility_criteria_option}")
        dispatcher.utter_message(text="Please answer the following questions:")
        return [SlotSet("indicator", indicator),
                FollowupAction("action_set_eligibility_criteria_question")]
    
class ActionEligibilityCriteriaQuestion(Action):
    def name(self) -> str:
        return "action_set_eligibility_criteria_question"
    

    def run(self, dispatcher: CollectingDispatcher, tracker: Tracker, domain: Dict[Text, Any]) -> List[Dict[Text, Any]]:
        question_list = eligibility_criteria_question[0]
        print(question_list)

        index = int(tracker.get_slot("question_index"))

        print(f"set question index - {index}")
        print(tracker.get_slot("eligibility_criteria_response"))

        if index >= len(question_list):
            dispatcher.utter_message(text="Completed..!!")
            return [ FollowupAction("action_store_response"),
                    SlotSet("question_index", 0)]
        else:
            dispatcher.utter_message(text=question_list[index])
            return[SlotSet("question0", question_list[index])]
    
class ActionValidateQuestionForm(FormValidationAction):
    def name(self):
        return "action_validate_eligibility_criteria"
    
    def run(self, dispatcher: CollectingDispatcher, tracker: Tracker, domain: Dict[Text, Any]) -> List[Dict[Text, Any]]:
        current_question = tracker.get_slot("question0")
        user_answer = tracker.latest_message.get('text')
        index = int(tracker.get_slot("question_index"))
        response_list = tracker.get_slot("eligibility_criteria_response")
        # print(response_list, "Q:" + current_question, "Ans" + user_answer)

        if response_list:
            updated_response_list = response_list + [user_answer]
        else:
            updated_response_list = user_answer
        
        prompt = yes_no_prompt.format(current_question=current_question, user_answer=user_answer)
        
        try:
            answer_relevance = prompt_engineering(prompt=prompt).lower()
        except Exception as e:
            print(f"Error occurred: {e}")
            
        print(answer_relevance)

        if answer_relevance == "yes.":
            updated_index = index + 1
            # print(updated_response_list)
            return[SlotSet("question_index", updated_index),
                   SlotSet("eligibility_criteria_response", updated_response_list),
                   FollowupAction("action_set_eligibility_criteria_question")]
            # return {"eligibility_criteria_response": updated_response_list}
        elif answer_relevance == "no.":
            dispatcher.utter_message(text="Please enter answer relevant to the question.")
            # return {"question_index": updated_index}
            return[FollowupAction("action_set_eligibility_criteria_question")]
    
##################################################################################################################

class ActionStoreResponse(Action):
    def name(self):
        return "action_store_response"
    
    def save_pdf(self, flexible_response, applied_context_response, eligibility_criteria_response):
        flexible_question = flexible_work_questions[0]
        applied_question = applied_content_questions[0]
        eligibility_question = eligibility_criteria_question[0]
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", style='B', size=14)
        # flexible work policy
        pdf.cell(200, 10, txt="Policy Questions and Responses", ln=True, align='C')
        pdf.ln(10)
        pdf.cell(200, 10, txt="1. Elements to be include in your flexible work policy?", ln=True, align='L')
        pdf.set_font("Arial", size=12)
        for num in range(len(flexible_question)):
            pdf.multi_cell(200, 10, txt="Q." + flexible_question[num], align='L')
            pdf.multi_cell(200, 10, txt= "Ans: " + flexible_response[num], align='L')
        
        pdf.ln(10)

        # contexts likely to apply flexible work policy
        pdf.set_font("Arial", style='B', size=14)
        pdf.cell(200, 10, txt="2. Contexts likely to apply on flexible work policy", ln=True, align='L')
        pdf.set_font("Arial", size=12)
        for num in range(len(applied_question)):
            pdf.multi_cell(200, 10, txt="Q." + applied_question[num], align='L')
            pdf.multi_cell(200, 10, txt= "Ans: " + applied_context_response[num], align='L')

        pdf.ln(10)

        # eligibility criteria
        pdf.set_font("Arial", style='B', size=14)
        pdf.cell(200, 10, txt="3. Eligibility criteria", ln=True, align='L')
        pdf.set_font("Arial", size=12)
        for num in range(len(eligibility_question)):
            pdf.multi_cell(200, 10, txt="Q." + eligibility_question[num], align='L')
            pdf.multi_cell(200, 10, txt= "Ans: " + eligibility_criteria_response[num], align='L')
        pdf.output("Policy_Questions_and_Responses.pdf")
         
    def run(self, dispatcher: CollectingDispatcher, tracker: Tracker, domain: dict):
        dispatcher.utter_message(text="Your responses are recorded")

        flexible_response = tracker.get_slot("user_response")
        applied_context_response = tracker.get_slot("applied_context_response")
        eligibility_criteria_response = tracker.get_slot("eligibility_criteria_response")
        self.save_pdf(flexible_response, applied_context_response, eligibility_criteria_response)

        dispatcher.utter_message(text="Saved in 'Policy_Questions_and_Responses.pdf' file!!!")
        return [FollowupAction("action_hr_policy_observations")] #[FollowupAction("action_select_applied_context")]
    
class actionCustomFallback(Action):
    def name(self):
        return "action_custom_fallback"       
        
    def run(self, dispatcher: CollectingDispatcher, tracker: Tracker, domain: Dict[Text, Any])  -> List[Dict[Text, Any]]:
        indicator = tracker.get_slot("indicator")
        print(indicator)

        action_map = {
            "flexible_work": "action_validate_flexible_work",
            "applied_context": "action_validate_applied_context_form",
            "eligibility_criteria": "action_validate_eligibility_criteria",
            "company_name": "action_set_company_name",
            "logo": "action_brand_color",
            "missing_element": "action_validate_missing_element_question"
        }
        # print(f"follow up action - {action_map[indicator]}")

        return [FollowupAction(action_map[indicator])]

        # if (indicator == "flexible_work"):
        #     return [FollowupAction("action_validate_flexible_work")]
        # elif (indicator == "applied_context"):
        #     return [FollowupAction("action_validate_applied_context_form")]
        # elif (indicator == "eligibility_criteria"):
        #     return [FollowupAction("action_validate_eligibility_criteria")]
        # elif (indicator == "custom_policy"):
        #     return [FollowupAction("action_track_custom_policy_input")]

#########################################################################################################################3

custom_policy_prompt = """

"""
class ActionGenerateCustomPolicyClause(Action):
    def name(self):
        return "action_generate_custom_policy_clause"
    
    def run(self, dispatcher: CollectingDispatcher, tracker: Tracker, domain: Dict[Text, Any])  -> List[Dict[Text, Any]]:
        dispatcher.utter_message(text="what would you like to add in your policy?")
        return[SlotSet("indicator", "custom_policy")
            # FollowupAction("action_track_custom_policy_input")
            ]
    
    class ActionTrackCustomPolicyInput(Action):
        def name(self):
            return "action_track_custom_policy_input"
        
        def run(self, dispatcher: CollectingDispatcher, tracker: Tracker, domain: Dict[Text, Any])  -> List[Dict[Text, Any]]:
            user_custom_policy = tracker.latest_message.get("text")
            # prompt = PROMPT_TEMPLATE.format(custom_policy=user_custom_policy)
            custom_policy = "custom_policy" #prompt_engineering(prompt=prompt)
            dispatcher.utter_message(text=f" Custom policy generated according to your instruction: {custom_policy}")
            return [SlotSet("custom_policy_input", custom_policy),
                    FollowupAction("action_hr_policy_observations")]
    
ai_observation_prompt_template = """
This is a remote work HR policy, are we missing something in this HR policy?
Identify the missing elements in the following remote work HR policy:

{policy}

If there are any missing element in the provided policy, Return the missing elements separated by commas.
"""

missing_element_prompt = """
"Generate questions that address the following missing elements in a Remote Work HR policy: {missing_elements}
Avoid compound and tail questions under any circumstances. 
Do not provide any additional information beyond the questions"

Separator: |
"""


missing_elements = []

class ActionHRPolicyObservation(Action):
    def name(self):
        return "action_hr_policy_observations"
    
    def extract_text_from_pdf(self, pdf_file_path):
        try:
            with open(pdf_file_path, 'rb') as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                text = ''
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text()
                return text
        except Exception as e:
            print(f"Error extracting text from PDF: {e}")
            return None
    
    def list_policy_elements(policy_string):
        policy_string = policy_string.replace(", and", ",")
        elements_list = [element.strip() for element in policy_string.split(',') if element.strip()]
        return elements_list

    def run(self, dispatcher: CollectingDispatcher, tracker: Tracker, domain: Dict[Text, Any])  -> List[Dict[Text, Any]]:
        pdf_file_path = "Policy_Questions_and_Responses.pdf"
        extracted_text = self.extract_text_from_pdf(pdf_file_path)

        prompt = ai_observation_prompt_template.format(policy = extracted_text)
        missing_element = prompt_engineering(prompt=prompt)
        ele = missing_element.replace(", and", ",")
        elements_list = [element.strip() for element in ele.split(',') if element.strip()]
        missing_elements.append(elements_list)
        dispatcher.utter_message(text=f"Missing element in your policy are:\n{missing_element}")
        return [FollowupAction("action_handle_missing_element")]
    
missing_element_questions = []

class ActionHandleMissingElement(Action):
    def name(self):
        return "action_handle_missing_element"
    
    def run(self, dispatcher: CollectingDispatcher, tracker: Tracker, domain: Dict[Text, Any])  -> List[Dict[Text, Any]]:
        elements = missing_elements[0]
        prompt = missing_element_prompt.format(missing_elements=elements)
        questions_list = prompt_engineering(prompt=prompt).replace("\n", "").replace("- ", "")

        questions = [question.strip() for question in questions_list.split('|')]
        print(questions)
        # for num in range(len(questions)):
        #     # if (num != 0) or (num % 3 != 0): # because of missing element response only 
        #             missing_element_questions.append(questions[num])
        missing_element_questions.append(questions)
        dispatcher.utter_message(text="\nPlease answer the following question to fill up the missing gap in your HR policy...\n")
        return [SlotSet("indicator", "missing_element"),
                FollowupAction("action_ask_question_missing_element")]
    
class ActionAskQMiuestoinssingElement(Action):
    def name(self):
        return "action_ask_question_missing_element"
    
    def run(self, dispatcher: CollectingDispatcher, tracker: Tracker, domain: Dict[Text, Any])  -> List[Dict[Text, Any]]:
        question_list = missing_element_questions[0]
        print(question_list)
        # if (num != 0) or (num % 3 != 0):
        index = int(tracker.get_slot("question_index"))
        # if (index == 0) or (index % 3 == 0):
        #     index = index+1
        print(f"set question index - {index}")
        print(tracker.get_slot("missing_element_response"))

        if index >= len(question_list):         
            # dispatcher.utter_message(text="Completed..!!")
            return [ FollowupAction("action_logo"),
                    SlotSet("question_index", 0)]
        else:
            dispatcher.utter_message(text=question_list[index])
            return[SlotSet("question0", question_list[index])]
    
class ActionValidateMissingElementQuestion(Action):
    def name(self):
        return "action_validate_missing_element_question"
    
    def run(self, dispatcher: CollectingDispatcher, tracker: Tracker, domain: Dict[Text, Any])  -> List[Dict[Text, Any]]:
        current_question = tracker.get_slot("question0")
        user_answer = tracker.latest_message.get('text')
        index = int(tracker.get_slot("question_index"))
        response_list = tracker.get_slot("missing_element_response")
        # print(response_list, "Q:" + current_question, "Ans" + user_answer)

        if response_list:
            updated_response_list = response_list + [user_answer]
        else:
            updated_response_list = user_answer
        
        prompt = yes_no_prompt.format(current_question=current_question, user_answer=user_answer)
        
        try:
            answer_relevance = prompt_engineering(prompt=prompt).lower()
        except Exception as e:
            print(f"Error occurred: {e}")
            
        print(answer_relevance)

        if answer_relevance == "yes.":
            updated_index = index + 1
            # print(updated_response_list)
            return[SlotSet("question_index", updated_index),
                   SlotSet("missing_element_response", updated_response_list),
                   FollowupAction("action_ask_question_missing_element")]
            # return {"eligibility_criteria_response": updated_response_list}
        elif answer_relevance == "no.":
            dispatcher.utter_message(text="Please enter answer relevant to the question.")
            # return {"question_index": updated_index}
            return[FollowupAction("action_ask_question_missing_element")]

        return[FollowupAction("action_ask_question_missing_element")]
    
class ActionHRPolicyObservation(Action):
    def name(self):
        return "action_handle_policy_feedback"
    
    def run(self, dispatcher: CollectingDispatcher, tracker: Tracker, domain: Dict[Text, Any])  -> List[Dict[Text, Any]]:
        user_response = tracker.latest_message.get("text").lower()

        if "yes" in user_response:
            dispatcher.utter_message(text="working on your feeedback...")
        elif "no" in user_response:
            dispatcher.utter_message(text="Great! Can you provide the company's logo and brand color for inclusion in the policy?")
            return [FollowupAction("action_logo_and_brand_color")]

class ActionMissingPolicy(Action):
    def name(self):
        return "action_missing_policy"
    
    def run(self, dispatcher: CollectingDispatcher, tracker: Tracker, domain: Dict[Text, Any])  -> List[Dict[Text, Any]]:
        dispatcher.utter_message(text="Handled missing part")
        return[FollowupAction("action_logo_and_brand_color")]

class ActionLogoBrandColor(Action):
    def name(self):
        return "action_logo" #and_brand_color"
    
    def run(self, dispatcher: CollectingDispatcher, tracker: Tracker, domain: Dict[Text, Any])  -> List[Dict[Text, Any]]:
        dispatcher.utter_message(text="Please provide url of brand logo")
        return [SlotSet("indicator", "logo")]
    
class ActionLogoBrandColor(Action):
    def name(self):
        return "action_brand_color"
    
    def run(self, dispatcher: CollectingDispatcher, tracker: Tracker, domain: Dict[Text, Any])  -> List[Dict[Text, Any]]:
        brand_logo = tracker.latest_message.get("text")#.lower()
        # dispatcher.utter_message(image=brand_logo)
        dispatcher.utter_message(text="Please provide color of brand")

        return [SlotSet("indicator", "color"),
                SlotSet("logo_url", brand_logo)]
    
        # return [FollowupAction("action_create_policy_document")]
    
class ActionConfirmBrandLogoColor(Action):
    def name(self):
        return "action_confirm_brand_logo_color"
    
    def run(self, dispatcher: CollectingDispatcher, tracker: Tracker, domain: Dict[Text, Any])  -> List[Dict[Text, Any]]:
        brand_color = tracker.latest_message.get("text")
        # brand_logo = tracker.get_slot("logo_url")
        # dispatcher.utter_message(image=brand_logo)
        dispatcher.utter_message(f"Brand Color: {brand_color}")
        return[SlotSet("brand_color", brand_color),
               FollowupAction("action_create_policy_document")]

    
class ActionCreatePolicyDocument(Action):
    def name(self):
        return "action_create_policy_document"
    
    
    def createDocument(self, company_name, brand_color, brand_logo, flexible_response, applied_response, eligibility_response, missing_response):
        flexible_question = flexible_work_questions[0]
        applied_question = applied_content_questions[0]
        eligibility_question = eligibility_criteria_question[0]
        missing_question = missing_element_questions[0]

        doc = Document()

        table = doc.add_table(rows=1, cols=1)
        table.autofit = False
        table.columns[0].width = Inches(10)
        cell = table.cell(0, 0)

        # Fill the cell with color
        # Use RGBColor for color specification
        cell._element.get_or_add_tcPr().append(
            OxmlElement('w:shd', {
                qn('w:fill'): brand_color
            })
        )

        

        # table = doc.add_table(rows=1, cols=2)

        # cell_logo = table.cell(0, 0)
        # paragraph_logo = cell_logo.add_paragraph()
        # run_logo = paragraph_logo.add_run()
        # run_logo.add_picture('pexels-helloaesthe-28056131.jpg', width=Inches(0.5), height=Inches(0.5))

        # cell_text = table.cell(0, 1)
        # paragraph_text = cell_text.add_paragraph()
        # run_text = paragraph_text.add_run(company_name)
        # run_text.font.name = 'Times New Roman'
        # run_text.font.size = Pt(14)
        # run_text.bold = True
        doc.add_paragraph("\n")
        doc.add_picture(brand_logo, width=Inches(0.5), height=Inches(0.5))
        doc.add_paragraph(company_name)

        # Add a title
        title = doc.add_heading('HR Policy', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.runs[0]
        run.font.color.rgb = RGBColor(0,0,0)
        
        policy = doc.add_heading('Policy Question Answers', level=2)
        policy.add_run("\n")
        run = policy.runs[0]
        run.font.color.rgb = RGBColor(0,0,0)
        run.font.name = 'Times New Roman' 
        run.font.size = Pt(14)
        policy.add_run("\n")

        flexiblePolicy = doc.add_paragraph()
        run = flexiblePolicy.add_run("Flexible work policy")
        run.bold = True
        flexiblePolicy.add_run("\n")
        for num in range(len(flexible_question)):
            run = flexiblePolicy.add_run(flexible_question[num])
            flexiblePolicy.add_run("\n")
            run = flexiblePolicy.add_run(f"Ans: {flexible_response[num]}")
            flexiblePolicy.add_run("\n")

        for run in flexiblePolicy.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

        appliedContext = doc.add_paragraph()
        run = appliedContext.add_run("Applied Context")
        run.bold = True
        appliedContext.add_run("\n")
        for num in range(len(applied_question)):
            run = appliedContext.add_run(applied_question[num])
            appliedContext.add_run("\n")
            run = appliedContext.add_run(f"Ans: {applied_response[num]}")
            appliedContext.add_run("\n")

        for run in appliedContext.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

        eligibilityCriteria = doc.add_paragraph()
        run = eligibilityCriteria.add_run("Eligibility Criteria")
        run.bold = True
        eligibilityCriteria.add_run("\n")
        for num in range(len(eligibility_question)):
            run = eligibilityCriteria.add_run(eligibility_question[num])
            eligibilityCriteria.add_run("\n")
            run = eligibilityCriteria.add_run(f"Ans: {eligibility_response[num]}")
            eligibilityCriteria.add_run("\n")

        for run in eligibilityCriteria.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

        missingElement = doc.add_paragraph()
        run = missingElement.add_run("Missing Element")
        run.bold = True
        missingElement.add_run("\n")
        for num in range(len(missing_question)):
            run = missingElement.add_run(missing_question[num])
            missingElement.add_run("\n")
            run = missingElement.add_run(f"Ans: {missing_response[num]}")
            missingElement.add_run("\n")

        for run in missingElement.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

        # Save the document
        doc.save(f'{company_name}_HR_document.docx')
        return "Document Created..."

    def run(self, dispatcher: CollectingDispatcher, tracker: Tracker, domain: Dict[Text, Any])  -> List[Dict[Text, Any]]:
        company_name = tracker.get_slot("company_name")
        brand_color = tracker.get_slot("brand_color")
        brand_logo = tracker.get_slot("logo_url")
        flexible_response = tracker.get_slot("user_response")
        applied_context_response = tracker.get_slot("applied_context_response")
        eligibility_criteria_response = tracker.get_slot("eligibility_criteria_response")
        missing_response = tracker.get_slot("missing_element_response")
        self.createDocument(company_name, brand_color,brand_logo, flexible_response, applied_context_response, eligibility_criteria_response, missing_response)
        dispatcher.utter_message(text="HR policy documnet created. Thank you for time!!")
        return[]
